#!/usr/bin/env python3
"""
Render VR (Vetenskapsrådet) publication list as Word (.docx).

Usage:
    pip install python-docx   # one-time
    python scripts/render_vr_docx.py

Output: output/vr_publist.docx

Reads cv_data.yaml and vr_config.yaml (same config as render_vr.py).
"""

import os
import re
import sys

# Force UTF-8 stdout on Windows
if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

import yaml

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def load_yaml(path):
    with open(path, encoding="utf-8") as f:
        content = f.read()
    return yaml.load(content, Loader=yaml.SafeLoader)


def format_authors_vancouver(authors_str, max_authors=6):
    if not authors_str:
        return ""
    authors_str = str(authors_str).strip()
    if "et al" in authors_str.lower() and authors_str.count(",") < 15:
        return authors_str

    parts = [p.strip() for p in authors_str.split(",")]
    two_element_format = False
    if len(parts) >= 4:
        initials_count = 0
        for j in range(1, min(6, len(parts)), 2):
            p = parts[j].strip()
            if (len(p) <= 4 and p.replace("-", "").replace(" ", "").replace(".", "").isalpha()
                    and p[0].isupper()):
                initials_count += 1
        if initials_count >= 2:
            two_element_format = True

    if two_element_format:
        authors = []
        i = 0
        while i < len(parts):
            if i + 1 < len(parts):
                name = parts[i].strip()
                initials = parts[i + 1].strip()
                if (len(initials) <= 4 and
                    initials.replace("-", "").replace(" ", "").replace(".", "").isalpha()):
                    authors.append(f"{name} {initials}")
                    i += 2
                    continue
            authors.append(parts[i].strip())
            i += 1
    else:
        authors = [p.strip() for p in parts if p.strip()]

    if len(authors) > max_authors:
        return ", ".join(authors[:6]) + ", et al"
    return ", ".join(authors)


def format_reference(pub):
    """Build a plain-text reference string (without bold name — that's handled by add_reference)."""
    parts = []
    if pub.get("authors"):
        parts.append(format_authors_vancouver(pub["authors"]))
        parts.append(". ")
    title = (pub.get("title", "") or "").rstrip(".")
    parts.append(title)
    parts.append(". ")
    if pub.get("journal"):
        parts.append(pub["journal"].rstrip("."))
    if pub.get("year"):
        parts.append(f" {pub['year']}")
    if pub.get("volume"):
        parts.append(f";{pub['volume']}")
    if pub.get("issue"):
        parts.append(f"({pub['issue']})")
    if pub.get("pages"):
        parts.append(f":{pub['pages']}")
    elif pub.get("article_number"):
        parts.append(f":{pub['article_number']}")
    parts.append(".")
    if pub.get("doi"):
        parts.append(f" doi: {pub['doi']}")
    return "".join(parts)


def add_reference(paragraph, pub, bold_name="Nilsonne G"):
    """Add a formatted reference to a paragraph, with the applicant name in bold
    and journal in italics."""
    authors = format_authors_vancouver(pub.get("authors", ""))
    title = (pub.get("title", "") or "").rstrip(".")
    doi = pub.get("doi", "")

    # Authors with bold name
    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                paragraph.add_run(before)
            paragraph.add_run(bold_name).bold = True
            if after:
                paragraph.add_run(after)
        else:
            paragraph.add_run(authors)
        paragraph.add_run(". ")

    # Title
    paragraph.add_run(title + ". ")

    # Journal (italic)
    journal = (pub.get("journal", "") or "").rstrip(".")
    if journal:
        run = paragraph.add_run(journal)
        run.italic = True

    # Year;Volume(Issue):Pages
    ref_detail = ""
    if pub.get("year"):
        ref_detail += f" {pub['year']}"
    if pub.get("volume"):
        ref_detail += f";{pub['volume']}"
    if pub.get("issue"):
        ref_detail += f"({pub['issue']})"
    if pub.get("pages"):
        ref_detail += f":{pub['pages']}"
    elif pub.get("article_number"):
        ref_detail += f":{pub['article_number']}"
    ref_detail += "."
    paragraph.add_run(ref_detail)

    # DOI
    if doi:
        paragraph.add_run(f" doi: {doi}")


def add_simple_reference(paragraph, pub, bold_name="Nilsonne G"):
    """Simplified reference for talks, preprints, etc."""
    authors = format_authors_vancouver(pub.get("authors", ""))
    title = (pub.get("title", "") or "").rstrip(".")

    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                paragraph.add_run(before)
            paragraph.add_run(bold_name).bold = True
            if after:
                paragraph.add_run(after)
        else:
            paragraph.add_run(authors)
        paragraph.add_run(". ")

    paragraph.add_run(title + ".")

    journal = (pub.get("journal", "") or "").rstrip(".")
    if journal:
        paragraph.add_run(" ")
        run = paragraph.add_run(journal)
        run.italic = True
        paragraph.add_run(".")

    if pub.get("year"):
        paragraph.add_run(f" {pub['year']}.")

    if pub.get("doi"):
        paragraph.add_run(f" doi: {pub['doi']}")


def in_year_range(pub, start=2017, end=2025):
    year = pub.get("year", 0) or 0
    return start <= year <= end


def classify_publication(pub, type_overrides=None):
    if type_overrides and pub.get("doi") in type_overrides:
        return type_overrides[pub["doi"]]
    vr_type = pub.get("vr_type")
    if vr_type:
        return vr_type
    title = (pub.get("title", "") or "").lower()
    if any(x in title for x in ["systematic review", "meta-analy", "scoping review"]):
        return "review"
    return "original"


def set_style(doc):
    """Configure document styles to match VR requirements."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Also set the east-asian and complex-script font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")

    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_numbered_item(doc, pub, number, bold_name="Nilsonne G", full=True):
    """Add a numbered reference item."""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if full:
        add_reference(p, pub, bold_name)
    else:
        add_simple_reference(p, pub, bold_name)
    return p


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base)

    data = load_yaml("cv_data.yaml")

    vr_config = {}
    if os.path.exists("vr_config.yaml"):
        vr_config = load_yaml("vr_config.yaml") or {}

    type_overrides = {}
    for item in vr_config.get("type_overrides", []) or []:
        if "doi" in item and "type" in item:
            type_overrides[item["doi"]] = item["type"]

    publications = data.get("publications", [])
    preprints = data.get("preprints", [])
    books = data.get("books", [])
    scholarly_debate = data.get("scholarly_debate", [])
    conference_abstracts = data.get("conference_abstracts", [])
    popular_science_writings = data.get("popular_science_writings", [])
    reports = data.get("reports", [])
    digital_research_objects = data.get("digital_research_objects", [])

    # === Prepare data (same logic as render_vr.py) ===

    # Section 1: Selected outputs
    selected_dois = vr_config.get("selected_outputs", []) or []
    contributions = {item["doi"]: item.get("contribution", "")
                     for item in (vr_config.get("selected_contributions", []) or [])
                     if "doi" in item}
    selected_outputs = []
    for doi in selected_dois:
        for pub in publications + preprints + books:
            if pub.get("doi") == doi:
                pub_copy = dict(pub)
                pub_copy["vr_contribution"] = contributions.get(doi, "[Describe contribution here]")
                selected_outputs.append(pub_copy)
                break

    # Section 2: Peer-reviewed 2017-2025
    pubs_sorted = sorted(publications, key=lambda p: -(p.get("year", 0) or 0))
    peer_original_2017 = []
    peer_reviews_2017 = []
    peer_conference_2017 = []
    peer_books_2017 = []
    peer_other_2017 = []

    for pub in pubs_sorted:
        if not in_year_range(pub):
            continue
        ptype = classify_publication(pub, type_overrides)
        if ptype == "review":
            peer_reviews_2017.append(pub)
        elif ptype == "conference":
            peer_conference_2017.append(pub)
        elif ptype == "book":
            peer_books_2017.append(pub)
        elif ptype == "other":
            peer_other_2017.append(pub)
        else:
            peer_original_2017.append(pub)

    for ca in sorted(conference_abstracts, key=lambda p: -(p.get("year", 0) or 0)):
        if in_year_range(ca):
            peer_conference_2017.append(ca)

    for b in books:
        if in_year_range(b):
            peer_books_2017.append(b)

    # Section 3: Non peer-reviewed 2017-2025
    popular_science_2017 = []
    for pw in sorted(popular_science_writings, key=lambda p: -(p.get("year", 0) or 0)):
        if in_year_range(pw):
            popular_science_2017.append(pw)
    for sd in sorted(scholarly_debate, key=lambda p: -(p.get("year", 0) or 0)):
        if in_year_range(sd):
            popular_science_2017.append(sd)

    preprints_2017 = sorted(
        [p for p in preprints if in_year_range(p)],
        key=lambda p: -(p.get("year", 0) or 0)
    )

    nonpeer_other_2017 = []
    for r in sorted(reports, key=lambda p: -(p.get("year", 0) or 0)):
        if in_year_range(r):
            nonpeer_other_2017.append(r)
    for dro in sorted(digital_research_objects, key=lambda p: -(p.get("year", 0) or 0)):
        if in_year_range(dro):
            nonpeer_other_2017.append(dro)

    # Section 4: Counts
    all_originals = [p for p in publications if classify_publication(p, type_overrides) == "original"]
    all_reviews = [p for p in publications if classify_publication(p, type_overrides) == "review"]
    all_other = len(books) + len(reports) + len(scholarly_debate) + len(popular_science_writings) + len(preprints)
    other_2017 = (
        len([b for b in books if in_year_range(b)]) +
        len([r for r in reports if in_year_range(r)]) +
        len([s for s in scholarly_debate if in_year_range(s)]) +
        len([p for p in popular_science_writings if in_year_range(p)]) +
        len(preprints_2017)
    )

    # === Build Word document ===
    doc = Document()
    set_style(doc)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{data.get('meta', {}).get('name', '')} – Publication List")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # === Section 1 ===
    add_heading(doc, "1. Selection of research outputs", level=2)
    p = doc.add_paragraph()
    run = p.add_run("The 10 publications most important for confirming competence as project leader and for the proposed research.")
    run.italic = True
    run.font.size = Pt(10)

    if selected_outputs:
        for i, pub in enumerate(selected_outputs, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.add_run(f"{i}. ").bold = True
            add_reference(p, pub)
            # Contribution
            contrib = pub.get("vr_contribution", "")
            if contrib:
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Cm(0.5)
                cp.paragraph_format.space_after = Pt(2)
                run = cp.add_run(contrib)
                run.italic = True
                run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("[Configure selected_outputs in vr_config.yaml]")
        run.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)

    # === Section 2 ===
    add_heading(doc, "2. Relevant peer-reviewed research outputs 2017–2025", level=2)

    sections_2 = [
        ("Original articles", peer_original_2017, True),
        ("Conference contributions", peer_conference_2017, False),
        ("Research review articles", peer_reviews_2017, True),
        ("Books and book chapters", peer_books_2017, False),
        ("Other outputs", peer_other_2017, False),
    ]

    for heading, items, full_ref in sections_2:
        if not items:
            continue
        add_heading(doc, heading, level=3)
        for i, pub in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.add_run(f"{i}. ").bold = False
            if full_ref:
                add_reference(p, pub)
            else:
                add_simple_reference(p, pub)

    # === Section 3 ===
    add_heading(doc, "3. Relevant non peer-reviewed research outputs 2017–2025", level=2)

    sections_3 = [
        ("Publications including popular science books/presentations", popular_science_2017),
        ("Preprints", preprints_2017),
        ("Other outputs", nonpeer_other_2017),
    ]

    for heading, items in sections_3:
        if not items:
            continue
        add_heading(doc, heading, level=3)
        for i, pub in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.add_run(f"{i}. ")
            add_simple_reference(p, pub)

    # === Section 4 ===
    add_heading(doc, "4. Number of publications", level=2)

    counts = [
        (f"Total number of peer-reviewed original articles", len(all_originals)),
        (f"Total number of peer-reviewed research review articles", len(all_reviews)),
        (f"Total number of other publications including patents", all_other),
        (f"Number of peer-reviewed original articles 2017–2025",
         len([p for p in all_originals if in_year_range(p)])),
        (f"Number of peer-reviewed research review articles 2017–2025",
         len([p for p in all_reviews if in_year_range(p)])),
        (f"Number of other publications including patents 2017–2025", other_2017),
    ]

    for label, count in counts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f"{label}: ").bold = False
        p.add_run(str(count)).bold = True

    # Save
    os.makedirs("output", exist_ok=True)
    out_path = os.path.join("output", "vr_publist.docx")
    doc.save(out_path)
    print(f"Written: {out_path}")

    # Summary
    print(f"\n=== VR Publication List Summary ===")
    print(f"Section 1 - Selected outputs: {len(selected_outputs)}")
    print(f"Section 2 - Peer-reviewed 2017-2025:")
    print(f"  Original articles: {len(peer_original_2017)}")
    print(f"  Conference contributions: {len(peer_conference_2017)}")
    print(f"  Research reviews: {len(peer_reviews_2017)}")
    print(f"  Books: {len(peer_books_2017)}")
    print(f"  Other: {len(peer_other_2017)}")
    print(f"Section 3 - Non peer-reviewed 2017-2025:")
    print(f"  Popular science: {len(popular_science_2017)}")
    print(f"  Preprints: {len(preprints_2017)}")
    print(f"  Other: {len(nonpeer_other_2017)}")
    print(f"Section 4 - Counts:")
    for label, count in counts:
        print(f"  {label}: {count}")


if __name__ == "__main__":
    main()
