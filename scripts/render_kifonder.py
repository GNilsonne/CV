#!/usr/bin/env python3
"""
Render KI fonder publication materials from cv_data.yaml.

Produces:
  - output/kifonder_selected_publications.docx
  - output/kifonder_publications_2021_present.pdf
  - output/kifonder_publications_2021_present.tex

Configuration:
  kifonder_config.yaml

Requirements:
  pip install python-docx pyyaml jinja2
  pdflatex available on PATH for PDF generation

Usage:
  python scripts/render_kifonder.py
  python scripts/render_kifonder.py --no-pdf
"""

import argparse
import os
import subprocess
import sys

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

import yaml
from jinja2 import Environment, FileSystemLoader

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx pyyaml jinja2", file=sys.stderr)
    sys.exit(1)


def load_yaml(path):
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f.read()) or {}


def add_hyperlink(paragraph, text, url, font_size=Pt(11), color=RGBColor(0, 0, 238)):
    """Add a clickable hyperlink to a paragraph in a Word document."""
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run_elem = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), str(color) if color else "0000EE")
    r_pr.append(color_elem)

    underline_elem = OxmlElement("w:u")
    underline_elem.set(qn("w:val"), "single")
    r_pr.append(underline_elem)

    if font_size:
        size_elem = OxmlElement("w:sz")
        size_elem.set(qn("w:val"), str(int(font_size.pt * 2)))
        r_pr.append(size_elem)

    run_elem.append(r_pr)
    run_elem.text = text
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def format_authors_vancouver(authors_str, max_authors=None):
    """Vancouver-style author formatting. max_authors=None means list all authors."""
    if not authors_str:
        return ""

    authors_str = str(authors_str).strip()
    parts = [p.strip() for p in authors_str.split(",")]

    two_element_format = False
    if len(parts) >= 4:
        initials_count = 0
        for j in range(1, min(6, len(parts)), 2):
            part = parts[j].strip()
            cleaned = part.replace("-", "").replace(" ", "").replace(".", "")
            if len(part) <= 4 and cleaned.isalpha() and part[:1].isupper():
                initials_count += 1
        if initials_count >= 2:
            two_element_format = True

    if two_element_format:
        authors = []
        i = 0
        while i < len(parts):
            if i + 1 < len(parts):
                surname = parts[i].strip()
                initials = parts[i + 1].strip()
                cleaned = initials.replace("-", "").replace(" ", "").replace(".", "")
                if len(initials) <= 4 and cleaned.isalpha():
                    authors.append(f"{surname} {initials}")
                    i += 2
                    continue
            authors.append(parts[i].strip())
            i += 1
    else:
        authors = [p.strip() for p in parts if p.strip()]

    while authors and authors[-1].lower().rstrip(".") in ("et al", ""):
        authors.pop()

    if max_authors and len(authors) > max_authors:
        return ", ".join(authors[:max_authors]) + ", et al"
    return ", ".join(authors)


def classify_publication(pub, type_overrides=None):
    """Classify as 'original', 'review', or 'other'."""
    pub_id = pub.get("id", "")
    if type_overrides:
        for item in type_overrides:
            if item.get("id") == pub_id and pub_id:
                return item["type"]

    if pub.get("pub_type"):
        return pub["pub_type"]

    title = (pub.get("title", "") or "").lower()
    if any(marker in title for marker in ["systematic review", "meta-analy", "scoping review"]):
        return "review"
    return "original"


def in_year_range(pub, start=2021, end=9999):
    year = pub.get("year", 0) or 0
    return start <= year <= end


def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")

    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_reference_para(doc, pub, number, bold_name="Nilsonne G"):
    """Cancerfonden-style numbered reference with bolded applicant name."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.space_before = Pt(1)

    indent = Cm(1.0)
    paragraph.paragraph_format.left_indent = indent
    paragraph.paragraph_format.first_line_indent = -indent

    number_run = paragraph.add_run(f"{number}.\t")
    number_run.font.size = Pt(11)

    p_pr = paragraph._element.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = p_pr.makeelement(qn("w:tabs"), {})
        p_pr.append(tabs)
    tab = tabs.makeelement(qn("w:tab"), {
        qn("w:val"): "left",
        qn("w:pos"): str(int(indent.emu / 914400 * 1440)),
    })
    tabs.append(tab)

    authors = format_authors_vancouver(pub.get("authors", ""))
    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                paragraph.add_run(before)
            paragraph.add_run(bold_name).bold = True
            if after:
                paragraph.add_run(after)
        else:
            paragraph.add_run(authors)
        paragraph.add_run(": ")

    title = (pub.get("title", "") or "").rstrip(".")
    paragraph.add_run(title)
    paragraph.add_run(". ")

    journal = (pub.get("journal", "") or pub.get("canonical_journal", "") or "").rstrip(".")
    if journal:
        journal_run = paragraph.add_run(journal)
        journal_run.italic = True

    detail = ""
    if pub.get("volume"):
        detail += f",{pub['volume']}"
        if pub.get("pages"):
            detail += f":{pub['pages']}"
        elif pub.get("article_number"):
            detail += f":{pub['article_number']}"
    if pub.get("year"):
        detail += f", {pub['year']}"
    detail += "."
    paragraph.add_run(detail)

    doi = (pub.get("doi", "") or "").strip()
    if doi:
        paragraph.add_run(" ")
        add_hyperlink(paragraph, f"doi:{doi}", f"https://doi.org/{doi}", font_size=Pt(11))


def tex_escape(text):
    if text is None:
        return ""
    text = str(text)
    replacements = {
        "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def tex_format_authors(authors_str):
    return tex_escape(format_authors_vancouver(authors_str, max_authors=None))


def tex_bold_name(text, name="Nilsonne G"):
    if not text:
        return ""
    escaped_name = tex_escape(name)
    if escaped_name in text:
        return text.replace(escaped_name, r"\textbf{" + escaped_name + "}")
    return text


def build_publication_index(publications):
    index = {}
    for pub in publications:
        pub_id = (pub.get("id", "") or "").strip()
        if pub_id:
            index[pub_id] = pub
    return index


def resolve_selected_publications(config_ids, publications):
    if len(config_ids) != 10:
        raise ValueError(
            f"kifonder_config.yaml must contain exactly 10 ids under selected_publications; found {len(config_ids)}."
        )

    publication_index = build_publication_index(publications)
    selected = []

    for pub_id in config_ids:
        if pub_id not in publication_index:
            raise ValueError(f"Selected publication id not found: {pub_id}")
        selected.append(publication_index[pub_id])

    return selected


def render_selected_docx(applicant_name, selected_publications, bold_name):
    doc = Document()
    set_style(doc)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{applicant_name} – Selected publications")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.name = "Times New Roman"

    add_heading(doc, "10 selected publications", level=2)
    for idx, pub in enumerate(selected_publications, start=1):
        add_reference_para(doc, pub, idx, bold_name=bold_name)

    os.makedirs("output", exist_ok=True)
    out_path = os.path.join("output", "kifonder_selected_publications.docx")
    doc.save(out_path)
    return out_path


def render_full_publication_pdf(data, config, publications):
    env = Environment(
        loader=FileSystemLoader("templates"),
        block_start_string="<%",
        block_end_string="%>",
        variable_start_string="<<",
        variable_end_string=">>",
        comment_start_string="<#",
        comment_end_string="#>",
    )
    env.filters["tex"] = tex_escape
    env.filters["notrailingdot"] = lambda s: str(s).rstrip(".") if s else ""
    env.filters["kif_authors"] = tex_format_authors
    env.filters["kif_bold_name"] = lambda s: tex_bold_name(s, config.get("bold_name", "Nilsonne G"))

    type_overrides = config.get("type_overrides", []) or []
    start_year = int(config.get("full_list_start_year", 2021))

    pubs_sorted = sorted(publications, key=lambda p: ((p.get("year", 0) or 0), p.get("date", "")), reverse=True)

    full_publication_list = []
    for pub in pubs_sorted:
        if not in_year_range(pub, start=start_year):
            continue
        ptype = classify_publication(pub, type_overrides)
        if ptype == "other":
            continue
        full_publication_list.append(pub)

    template = env.get_template("kifonder_publist.tex.j2")
    tex_output = template.render(
        meta=data.get("meta", {}),
        start_year=start_year,
        full_publication_list=full_publication_list,
        bold_name=config.get("bold_name", "Nilsonne G"),
    )

    os.makedirs("output", exist_ok=True)
    tex_path = os.path.join("output", "kifonder_publications_2021_present.tex")
    with open(tex_path, "w", encoding="utf-8") as handle:
        handle.write(tex_output)

    return tex_path


def compile_pdf(tex_path):
    workdir = os.path.dirname(tex_path)
    filename = os.path.basename(tex_path)
    commands = [
        ["pdflatex", "-interaction=nonstopmode", filename],
        ["pdflatex", "-interaction=nonstopmode", filename],
    ]
    for cmd in commands:
        result = subprocess.run(cmd, cwd=workdir, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"pdflatex failed for {' '.join(cmd)}\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
            )


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--no-pdf", action="store_true", help="Write the .tex file but skip PDF compilation")
    args = parser.parse_args()

    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base)

    data = load_yaml("cv_data.yaml")
    config = load_yaml("kifonder_config.yaml")

    publications = data.get("publications", []) or []
    selected_ids = config.get("selected_publications", []) or []
    applicant_name = data.get("meta", {}).get("name", "") or ""
    bold_name = config.get("bold_name", "Nilsonne G")

    selected_publications = resolve_selected_publications(selected_ids, publications)

    docx_path = render_selected_docx(applicant_name, selected_publications, bold_name)
    tex_path = render_full_publication_pdf(data, config, publications)

    print(f"Written: {docx_path}")
    print(f"Written: {tex_path}")

    if not args.no_pdf:
        compile_pdf(tex_path)
        pdf_path = tex_path[:-4] + ".pdf"
        print(f"Written: {pdf_path}")


if __name__ == "__main__":
    main()
