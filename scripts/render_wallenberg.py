#!/usr/bin/env python3
"""
Render Marcus and Amalia Wallenberg Foundation CV materials from cv_data.yaml.

Produces two Word documents:
  - output/wallenberg_cv.docx              (CV with initial sections only,
                                             no publications or later sections)
  - output/wallenberg_selected_pubs.docx   (10 selected publications)

Configuration in wallenberg_config.yaml:
  selected_publications:   list of publication IDs from cv_data.yaml

Usage:
    pip install python-docx   # one-time
    python scripts/render_wallenberg.py
"""

import os
import sys

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

import yaml

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx",
          file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_yaml(path):
    with open(path, encoding="utf-8") as f:
        return yaml.load(f.read(), Loader=yaml.SafeLoader)


def add_hyperlink(paragraph, text, url, font_size=Pt(11),
                  color=RGBColor(0, 0, 238)):
    """Insert a clickable hyperlink run into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000EE")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    run_elem.append(rPr)
    run_elem.text = text
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def format_authors_vancouver(authors_str, max_authors=6):
    """Vancouver-style: up to *max_authors*, then 'et al.'."""
    if not authors_str:
        return ""
    authors_str = str(authors_str).strip()
    if "et al" in authors_str.lower() and authors_str.count(",") < 15:
        return authors_str

    parts = [p.strip() for p in authors_str.split(",")]

    two_element_format = False
    if len(parts) >= 4:
        initials_count = 0
        for j in range(1, min(6, len(parts)), 2):
            p = parts[j].strip()
            if (len(p) <= 4
                    and p.replace("-", "").replace(" ", "").replace(".", "").isalpha()
                    and p[0].isupper()):
                initials_count += 1
        if initials_count >= 2:
            two_element_format = True

    if two_element_format:
        authors = []
        i = 0
        while i < len(parts):
            if i + 1 < len(parts):
                name = parts[i].strip()
                initials = parts[i + 1].strip()
                if (len(initials) <= 4
                        and initials.replace("-", "").replace(" ", "")
                                    .replace(".", "").isalpha()):
                    authors.append(f"{name} {initials}")
                    i += 2
                    continue
            authors.append(parts[i].strip())
            i += 1
    else:
        authors = [p.strip() for p in parts if p.strip()]

    if max_authors and len(authors) > max_authors:
        return ", ".join(authors[:max_authors]) + ", et al"
    return ", ".join(authors)


# ---------------------------------------------------------------------------
# Word styling helpers
# ---------------------------------------------------------------------------

def set_style(doc, font_name="Times New Roman", font_size=Pt(11)):
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = font_size

    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font_name)

    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_bullet(doc, text, bold=False):
    """Add a simple bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet_with_link(doc, text, url=None, link_label="[link]"):
    """Bullet with optional trailing hyperlink."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.add_run(text)
    if url:
        p.add_run(" ")
        add_hyperlink(p, link_label, url)
    return p


def add_sub_bullet(doc, text):
    """Add a second-level bullet (List Bullet 2)."""
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


# ---------------------------------------------------------------------------
# Reference formatting (for selected-publications document)
# ---------------------------------------------------------------------------

def add_reference_para(doc, pub, number, bold_name="Nilsonne G"):
    """Numbered reference with hanging indent, bold applicant name,
    italic journal, and DOI hyperlink."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(1)

    indent = Cm(1.0)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = -indent

    # Number
    num_run = p.add_run(f"{number}.\t")
    num_run.font.size = Pt(11)

    # Tab stop at indent
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = pPr.makeelement(qn("w:tabs"), {})
        pPr.append(tabs)
    tab = tabs.makeelement(qn("w:tab"), {
        qn("w:val"): "left",
        qn("w:pos"): str(int(indent.emu / 914400 * 1440)),
    })
    tabs.append(tab)

    # Authors (bold applicant name)
    authors = format_authors_vancouver(pub.get("authors", ""))
    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                p.add_run(before)
            p.add_run(bold_name).bold = True
            if after:
                p.add_run(after)
        else:
            p.add_run(authors)
        p.add_run(". ")

    # Title
    title = (pub.get("title", "") or "").rstrip(".")
    p.add_run(title + ". ")

    # Journal (italic)
    journal = (pub.get("journal", "") or "").rstrip(".")
    if journal:
        run = p.add_run(journal)
        run.italic = True

    # Volume(Issue):Pages, Year
    detail = ""
    if pub.get("year"):
        detail += f" {pub['year']}"
    if pub.get("volume"):
        detail += f";{pub['volume']}"
    if pub.get("issue"):
        detail += f"({pub['issue']})"
    if pub.get("pages"):
        detail += f":{pub['pages']}"
    elif pub.get("article_number"):
        detail += f":{pub['article_number']}"
    detail += "."
    p.add_run(detail)

    # DOI hyperlink
    doi = pub.get("doi", "") or ""
    if doi:
        p.add_run(" ")
        add_hyperlink(p, doi, f"https://doi.org/{doi}")

    return p


# ---------------------------------------------------------------------------
# Document 1: CV (initial sections only)
# ---------------------------------------------------------------------------

def build_cv(data, doc):
    """Populate *doc* with CV sections up to and including Other Commissions."""
    meta = data.get("meta", {})

    # --- Header ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(meta.get("name", ""))
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(meta.get("affiliation", ""))
    if meta.get("email"):
        sub.add_run(f"\n{meta['email']}")
    if meta.get("orcid"):
        sub.add_run(f"\nORCID: {meta['orcid']}")
    profiles = meta.get("profiles", {})
    if profiles:
        profile_parts = []
        if profiles.get("ki_profile"):
            profile_parts.append(("KI profile", profiles["ki_profile"]))
        # SU profile omitted — no longer current
        if profiles.get("google_scholar"):
            profile_parts.append(("Google Scholar", profiles["google_scholar"]))
        if profile_parts:
            link_p = doc.add_paragraph()
            link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, (label, url) in enumerate(profile_parts):
                if i > 0:
                    link_p.add_run("  ·  ")
                add_hyperlink(link_p, label, url)

    # --- Degrees ---
    degrees = data.get("degrees", [])
    if degrees:
        add_heading(doc, "Degrees and Qualifications", level=2)
        for deg in degrees:
            desc = deg.get("description", "")
            link = deg.get("link")
            add_bullet_with_link(doc, desc, url=link)

    # --- Employment ---
    employment = data.get("employment", [])
    if employment:
        add_heading(doc, "Employment and Appointments", level=2)
        for pos in employment:
            add_bullet(doc, pos.get("description", ""))

    # --- PhD Supervision ---
    supervision = data.get("phd_supervision", [])
    if supervision:
        add_heading(doc, "Supervision of PhD Students", level=2)
        for s in supervision:
            parts = [s.get("name", "")]
            if s.get("institution"):
                parts.append(s["institution"])
            parts.append(str(s.get("year", "")))
            if s.get("role"):
                parts.append(s["role"])
            line = ", ".join(p for p in parts if p)
            if s.get("thesis"):
                line += f'. Thesis: "{s["thesis"]}"'
            p = add_bullet(doc, line)
            if s.get("link"):
                p.add_run(" ")
                add_hyperlink(p, "[link]", s["link"])

    # --- Grants ---
    grants = data.get("grants", [])
    if grants:
        add_heading(doc, "Research Grants", level=2)
        for g in grants:
            desc = g.get("description", "")
            link = g.get("link")
            add_bullet_with_link(doc, desc, url=link)

    # --- Teaching ---
    teaching = data.get("teaching", [])
    if teaching:
        add_heading(doc, "Teaching", level=2)
        for t in teaching:
            add_bullet(doc, t.get("description", ""))
            for sub in t.get("subitems", []) or []:
                add_sub_bullet(doc, sub if isinstance(sub, str) else str(sub))

    # --- Awards ---
    awards = data.get("awards", [])
    if awards:
        add_heading(doc, "Awards", level=2)
        for a in awards:
            desc = a.get("description", "")
            link = a.get("link")
            add_bullet_with_link(doc, desc, url=link)

    # --- PhD Committee ---
    phd_committee = data.get("phd_committee", [])
    if phd_committee:
        add_heading(doc, "PhD Thesis Committee Member / Faculty Opponent", level=2)
        for c in phd_committee:
            add_bullet(doc, c.get("description", ""))

    # --- Academic Commissions ---
    academic_commissions = data.get("academic_commissions", [])
    if academic_commissions:
        add_heading(doc, "Academic Commissions of Trust", level=2)
        for c in academic_commissions:
            add_bullet(doc, c.get("description", ""))
            for sub in c.get("subitems", []) or []:
                if isinstance(sub, dict):
                    add_sub_bullet(doc, sub.get("name", str(sub)))
                else:
                    add_sub_bullet(doc, str(sub))

    # --- Other Commissions ---
    other_commissions = data.get("other_commissions", [])
    if other_commissions:
        add_heading(doc, "Other Commissions of Trust", level=2)
        for c in other_commissions:
            add_bullet(doc, c.get("description", ""))
            for sub in c.get("subitems", []) or []:
                add_sub_bullet(doc, sub if isinstance(sub, str) else str(sub))


# ---------------------------------------------------------------------------
# Document 2: Selected publications
# ---------------------------------------------------------------------------

def build_selected_pubs(data, config, doc):
    """Populate *doc* with the selected publications list."""
    meta = data.get("meta", {})

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(
        f"{meta.get('name', '')} – Selected Publications")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    selected_ids = config.get("selected_publications", []) or []
    all_pubs = (data.get("publications", [])
                + data.get("preprints", [])
                + data.get("books", []))

    selected = []
    for sel_id in selected_ids:
        for pub in all_pubs:
            if pub.get("id") == sel_id:
                selected.append(pub)
                break
        else:
            print(f"WARNING: selected publication id '{sel_id}' not found "
                  f"in cv_data.yaml")

    if selected:
        for i, pub in enumerate(selected, 1):
            add_reference_para(doc, pub, i)
    else:
        p = doc.add_paragraph()
        run = p.add_run(
            "[Configure selected_publications in wallenberg_config.yaml "
            "with up to 10 publication IDs]")
        run.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)

    return len(selected)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base)

    data = load_yaml("cv_data.yaml")

    config = {}
    if os.path.exists("wallenberg_config.yaml"):
        config = load_yaml("wallenberg_config.yaml") or {}

    os.makedirs("output", exist_ok=True)

    # --- Document 1: CV ---
    cv_doc = Document()
    set_style(cv_doc)
    build_cv(data, cv_doc)
    cv_path = os.path.join("output", "wallenberg_cv.docx")
    cv_doc.save(cv_path)
    print(f"Written: {cv_path}")

    # --- Document 2: Selected publications ---
    pub_doc = Document()
    set_style(pub_doc)
    n_selected = build_selected_pubs(data, config, pub_doc)
    pub_path = os.path.join("output", "wallenberg_selected_pubs.docx")
    pub_doc.save(pub_path)
    print(f"Written: {pub_path}")

    # Summary
    print(f"\n=== Wallenberg Output Summary ===")
    print(f"CV sections included: Degrees, Employment, PhD Supervision, "
          f"Grants, Teaching, Awards, PhD Committee, "
          f"Academic Commissions, Other Commissions")
    print(f"CV sections excluded: Publications and all subsequent sections")
    print(f"Selected publications: {n_selected}")


if __name__ == "__main__":
    main()
