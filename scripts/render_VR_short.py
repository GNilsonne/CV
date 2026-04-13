#!/usr/bin/env python3
"""
Render a short Vetenskapsradet publication list as Word (.docx).

Usage:
    pip install python-docx pyyaml   # one-time
    python scripts/render_VR_short.py

Output:
    output/vr_short.docx

Inputs:
    cv_data.yaml
    vr_short_config.yaml

Selection logic:
    - Select exactly 10 publications by id in vr_short_config.yaml.
    - If an id is missing, the script exits with an error.
"""

import os
import sys

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

import yaml

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Error: missing dependencies. Run: pip install python-docx pyyaml", file=sys.stderr)
    sys.exit(1)


def load_yaml(path):
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f.read())


def format_authors_vancouver(authors_str, max_authors=None):
    if not authors_str:
        return ""
    authors_str = str(authors_str).strip()

    parts = [p.strip() for p in authors_str.split(",")]

    two_element_format = False
    if len(parts) >= 4:
        initials_count = 0
        for j in range(1, min(6, len(parts)), 2):
            p = parts[j].strip()
            cleaned = p.replace("-", "").replace(" ", "").replace(".", "")
            if len(p) <= 4 and cleaned.isalpha() and p[0].isupper():
                initials_count += 1
        if initials_count >= 2:
            two_element_format = True

    if two_element_format:
        authors = []
        i = 0
        while i < len(parts):
            if i + 1 < len(parts):
                name = parts[i].strip()
                initials = parts[i + 1].strip()
                cleaned = initials.replace("-", "").replace(" ", "").replace(".", "")
                if len(initials) <= 4 and cleaned.isalpha():
                    authors.append(f"{name} {initials}")
                    i += 2
                    continue
            authors.append(parts[i].strip())
            i += 1
    else:
        authors = [p.strip() for p in parts if p.strip()]

    while authors and authors[-1].lower().rstrip(".") in ("et al", ""):
        authors.pop()

    if max_authors and len(authors) > max_authors:
        return ", ".join(authors[:max_authors]) + ", et al"
    return ", ".join(authors)


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:cs"), "Arial")

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def apply_run_style(run, *, bold=False, italic=False, size=11):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_title(doc, applicant_name):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(f"{applicant_name} - Selected publications")
    apply_run_style(run, bold=True, size=11)


def add_reference(doc, pub, number, bold_name="Nilsonne G"):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.75)

    run = paragraph.add_run(f"{number}. ")
    apply_run_style(run)

    authors = format_authors_vancouver(pub.get("authors", ""))
    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                apply_run_style(paragraph.add_run(before))
            apply_run_style(paragraph.add_run(bold_name), bold=True)
            if after:
                apply_run_style(paragraph.add_run(after))
        else:
            apply_run_style(paragraph.add_run(authors))
        apply_run_style(paragraph.add_run(". "))

    title = (pub.get("title", "") or "").rstrip(".")
    apply_run_style(paragraph.add_run(title + ". "))

    journal = (pub.get("journal", "") or pub.get("canonical_journal", "") or "").rstrip(".")
    if journal:
        apply_run_style(paragraph.add_run(journal), italic=True)

    details = ""
    if pub.get("year"):
        details += f" {pub['year']}"
    if pub.get("volume"):
        details += f";{pub['volume']}"
    if pub.get("issue"):
        details += f"({pub['issue']})"
    if pub.get("pages"):
        details += f":{pub['pages']}"
    elif pub.get("article_number"):
        details += f":{pub['article_number']}"
    if details:
        apply_run_style(paragraph.add_run(details))

    doi = (pub.get("doi", "") or "").strip()
    if doi:
        apply_run_style(paragraph.add_run(f". doi: {doi}"))
    else:
        apply_run_style(paragraph.add_run("."))


def build_publication_index(publications):
    index = {}
    for pub in publications:
        pub_id = (pub.get("id", "") or "").strip()
        if pub_id:
            index[pub_id] = pub
    return index


def resolve_selected_publications(config_ids, publications):
    if len(config_ids) != 10:
        raise ValueError(
            f"vr_short_config.yaml must contain exactly 10 ids under selected_publications; found {len(config_ids)}."
        )

    publication_index = build_publication_index(publications)
    selected = []

    for pub_id in config_ids:
        if pub_id not in publication_index:
            raise ValueError(f"Selected publication id not found: {pub_id}")
        selected.append(publication_index[pub_id])

    return selected


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base)

    data = load_yaml("cv_data.yaml") or {}
    config = load_yaml("vr_short_config.yaml") or {}

    publications = data.get("publications", []) or []
    selected_ids = config.get("selected_publications", []) or []
    applicant_name = data.get("meta", {}).get("name", "") or ""
    bold_name = config.get("bold_name", "Nilsonne G")

    selected_publications = resolve_selected_publications(selected_ids, publications)

    doc = Document()
    set_normal_style(doc)
    add_title(doc, applicant_name)

    for i, pub in enumerate(selected_publications, start=1):
        add_reference(doc, pub, i, bold_name=bold_name)

    os.makedirs("output", exist_ok=True)
    out_path = os.path.join("output", "vr_short.docx")
    doc.save(out_path)
    print(f"Written: {out_path}")
    print(f"Selected publications: {len(selected_publications)}")


if __name__ == "__main__":
    main()
