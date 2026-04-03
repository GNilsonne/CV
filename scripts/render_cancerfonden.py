#!/usr/bin/env python3
"""
Render Swedish Cancer Society (Cancerfonden) CV materials from cv_data.yaml.

Produces:
  - output/cancerfonden_publist.docx  (publication list PDF attachment)
  - output/cancerfonden_cv_fields.txt (text fields for copy-paste into application form)

Requirements from Cancerfonden:
  - Publication counts: total originals, first-author, last-author, reviews
  - Publication list PDF: 10 most important + complete list 2018-2026
    Three categories: peer-reviewed original, review articles, other articles
    Format: Authors (all), title, journal, volume, page, year. Name in bold.
  - Unpublished manuscripts (max 3, separate PDFs)
  - Other scientific merits (1500 chars)
  - PhD supervision: main completed, main ongoing, co completed, co ongoing (1000 chars each)

Usage:
    pip install python-docx   # one-time
    python scripts/render_cancerfonden.py
"""

import os
import sys
import textwrap

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

import yaml

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_hyperlink(paragraph, text, url, font_size=Pt(11), color=RGBColor(0, 0, 238)):
    """Add a clickable hyperlink to a paragraph in a Word document."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color) if color else "0000EE")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    run_elem.append(rPr)
    run_elem.text = text
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_yaml(path):
    with open(path, encoding="utf-8") as f:
        content = f.read()
    return yaml.load(content, Loader=yaml.SafeLoader)


def format_authors_vancouver(authors_str, max_authors=None):
    """Vancouver-style author formatting. max_authors=None means list ALL (Cancerfonden requirement)."""
    if not authors_str:
        return ""
    authors_str = str(authors_str).strip()

    parts = [p.strip() for p in authors_str.split(",")]

    # Detect two-element format: "Surname, I, Surname, I"
    two_element_format = False
    if len(parts) >= 4:
        initials_count = 0
        for j in range(1, min(6, len(parts)), 2):
            p = parts[j].strip()
            if (len(p) <= 4 and p.replace("-", "").replace(" ", "").replace(".", "").isalpha()
                    and p[0].isupper()):
                initials_count += 1
        if initials_count >= 2:
            two_element_format = True

    if two_element_format:
        authors = []
        i = 0
        while i < len(parts):
            if i + 1 < len(parts):
                name = parts[i].strip()
                initials = parts[i + 1].strip()
                if (len(initials) <= 4 and
                    initials.replace("-", "").replace(" ", "").replace(".", "").isalpha()):
                    authors.append(f"{name} {initials}")
                    i += 2
                    continue
            authors.append(parts[i].strip())
            i += 1
    else:
        authors = [p.strip() for p in parts if p.strip()]

    # Remove trailing "et al" — we want all authors
    while authors and authors[-1].lower().rstrip(".") in ("et al", ""):
        authors.pop()

    if max_authors and len(authors) > max_authors:
        return ", ".join(authors[:max_authors]) + ", et al"
    return ", ".join(authors)


def is_first_author(pub, name="Nilsonne"):
    """Check if applicant is first author."""
    authors = (pub.get("authors", "") or "").strip()
    return authors.startswith(name)


def is_last_author(pub, name="Nilsonne"):
    """Check if applicant is last author."""
    authors = (pub.get("authors", "") or "").strip()
    if not authors:
        return False
    segments = [s.strip() for s in authors.split(",")]
    # Skip trailing "et al"
    while segments and segments[-1].lower().rstrip(".") in ("et al", ""):
        segments.pop()
    if segments:
        return name in segments[-1]
    return False


def in_year_range(pub, start=2018, end=2026):
    year = pub.get("year", 0) or 0
    return start <= year <= end


def classify_publication(pub, type_overrides=None):
    """Classify as 'original', 'review', or 'other'.
    
    Priority: type_overrides (by id) > pub_type field > title heuristic.
    """
    pub_id = pub.get("id", "")
    if type_overrides:
        for item in type_overrides:
            if item.get("id") == pub_id and pub_id:
                return item["type"]
    # Use the pub_type field if present
    if pub.get("pub_type"):
        return pub["pub_type"]
    # Fallback heuristic
    title = (pub.get("title", "") or "").lower()
    if any(x in title for x in ["systematic review", "meta-analy", "scoping review"]):
        return "review"
    return "original"


def format_reference_text(pub):
    """Plain-text reference in Cancerfonden format:
    Authors (all), title, journal, volume, page, year.
    Example: Smith J, Doe A, Jane A: Analysis of cancer. J Can Res.,15:100-110, 2024
    """
    parts = []
    authors = format_authors_vancouver(pub.get("authors", ""))
    if authors:
        parts.append(authors)
        parts.append(": ")

    title = (pub.get("title", "") or "").rstrip(".")
    parts.append(title)
    parts.append(". ")

    journal = (pub.get("journal", "") or pub.get("canonical_journal", "") or "").rstrip(".")
    if journal:
        parts.append(journal)
        if pub.get("volume") or pub.get("year"):
            parts.append(",")

    if pub.get("volume"):
        parts.append(f"{pub['volume']}")
        if pub.get("pages"):
            parts.append(f":{pub['pages']}")
        elif pub.get("article_number"):
            parts.append(f":{pub['article_number']}")
        parts.append(",")

    if pub.get("year"):
        parts.append(f" {pub['year']}")

    # Clean up trailing commas/spaces
    result = "".join(parts).rstrip(",").rstrip()
    if not result.endswith("."):
        result += "."
    return result


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_reference_para(doc, pub, number, bold_name="Nilsonne G"):
    """Add a numbered reference with hanging indent: number in left margin,
    text block aligned. Applicant name bolded, journal italicised."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(1)

    # Hanging indent: left_indent sets the text block position,
    # first_line_indent (negative) pulls the number back into the margin.
    indent = Cm(1.0)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = -indent

    # Number (sits in the margin thanks to negative first-line indent)
    num_run = p.add_run(f"{number}.\t")
    num_run.font.size = Pt(11)

    # Set a tab stop at the indent position so \t aligns the text
    from docx.oxml.ns import qn as _qn
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(_qn("w:tabs"))
    if tabs is None:
        tabs = pPr.makeelement(_qn("w:tabs"), {})
        pPr.append(tabs)
    tab = tabs.makeelement(_qn("w:tab"), {
        _qn("w:val"): "left",
        _qn("w:pos"): str(int(indent.emu / 914400 * 1440)),  # convert to twips
    })
    tabs.append(tab)

    # Authors with bold name
    authors = format_authors_vancouver(pub.get("authors", ""))
    if authors:
        if bold_name in authors:
            before, after = authors.split(bold_name, 1)
            if before:
                p.add_run(before)
            p.add_run(bold_name).bold = True
            if after:
                p.add_run(after)
        else:
            p.add_run(authors)
        p.add_run(": ")

    # Title
    title = (pub.get("title", "") or "").rstrip(".")
    p.add_run(title + ". ")

    # Journal (italic)
    journal = (pub.get("journal", "") or "").rstrip(".")
    if journal:
        run = p.add_run(journal)
        run.italic = True

    # ,Volume:Pages, Year
    detail = ""
    if pub.get("volume"):
        detail += f",{pub['volume']}"
        if pub.get("pages"):
            detail += f":{pub['pages']}"
        elif pub.get("article_number"):
            detail += f":{pub['article_number']}"
    if pub.get("year"):
        detail += f", {pub['year']}"
    detail += "."
    p.add_run(detail)

    # DOI as hyperlink
    doi = pub.get("doi", "") or ""
    if doi:
        p.add_run(" ")
        doi_url = f"https://doi.org/{doi}"
        add_hyperlink(p, f"doi:{doi}", doi_url, font_size=Pt(11))

    # Open resource links: [preprint | data | code | ...]
    links = pub.get("links", {}) or {}
    if links:
        LINK_LABELS = {
            "preprint": "preprint", "data": "data", "code": "code",
            "materials": "materials", "preregistration": "preregistration",
            "narrative": "narrative", "slides": "slides", "video": "video",
            "protocol": "protocol", "web": "web", "poster": "poster",
            "correction": "correction", "pdf": "pdf", "diva": "DiVA",
            "program": "program", "url": "link",
        }
        link_parts = []
        seen_urls = set()
        for key, url in links.items():
            if not url or url in seen_urls:
                continue
            seen_urls.add(url)
            label = LINK_LABELS.get(key, key)
            link_parts.append((label, url))

        if link_parts:
            p.add_run(" [")
            for i, (label, url) in enumerate(link_parts):
                if i > 0:
                    p.add_run(" | ")
                add_hyperlink(p, label, url, font_size=Pt(11))
            p.add_run("]")

    return p


# ---------------------------------------------------------------------------
# Text field generators (for form copy-paste)
# ---------------------------------------------------------------------------

def generate_other_merits(data, max_chars=1500):
    """Auto-generate 'Other scientific merits' from YAML data."""
    lines = []

    # Invited lectures (count + recent highlights)
    talks = data.get("invited_talks", [])
    if talks:
        recent = [t for t in talks if (t.get("year", 0) or 0) >= 2020]
        lines.append(f"Invited lectures: {len(talks)} total ({len(recent)} since 2020).")

    # Academic commissions / leadership
    ac = data.get("academic_commissions", [])
    if ac:
        recent_ac = [a for a in ac if "20" in str(a.get("description", ""))[-10:]]
        # Pick a few highlights
        highlights = []
        for a in ac:
            desc = a.get("description", "")
            if any(kw in desc.lower() for kw in ["editor", "board", "chair", "director", "committee"]):
                highlights.append(desc)
        if highlights:
            lines.append("Leadership & editorial roles: " + "; ".join(highlights[:5]) + ".")

    # Awards
    awards = data.get("awards", [])
    if awards:
        lines.append("Awards: " + "; ".join(a.get("description", "") for a in awards) + ".")

    # Grants (count)
    grants = data.get("grants", [])
    if grants:
        recent_grants = [g for g in grants if any(str(y) in g.get("description", "") for y in range(2020, 2027))]
        lines.append(f"Research grants: {len(grants)} awarded ({len(recent_grants)} since 2020), including Horizon Europe, Riksbankens Jubileumsfond, Region Stockholm.")

    text = " ".join(lines)
    if len(text) > max_chars:
        text = text[:max_chars - 3] + "..."
    return text


def generate_supervision_fields(data):
    """Generate supervision text fields from phd_supervision YAML."""
    sup = data.get("phd_supervision", [])

    main_completed = []
    main_ongoing = []
    co_completed = []
    co_ongoing = []

    for s in sup:
        name = s.get("name", "")
        year = str(s.get("year", ""))
        role = (s.get("role", "") or "").lower()
        institution = s.get("institution", "")
        thesis = s.get("thesis", "")

        entry = f"{name}, {institution}"
        if thesis:
            entry += f", \"{thesis}\""

        ongoing = year.endswith("-") or "ongoing" in year.lower()

        if "main" in role or "principal" in role:
            if ongoing:
                planned = year.rstrip("-")
                main_ongoing.append(f"{entry}, planned {planned}.")
            else:
                main_completed.append(f"{entry}, {year}.")
        else:
            # co-supervisor
            if ongoing:
                planned = year.rstrip("-")
                co_ongoing.append(f"{entry}, planned {planned}.")
            else:
                co_completed.append(f"{entry}, {year}.")

    return {
        "main_completed": "\n".join(main_completed) if main_completed else "N/A",
        "main_ongoing": "\n".join(main_ongoing) if main_ongoing else "N/A",
        "co_completed": "\n".join(co_completed) if co_completed else "N/A",
        "co_ongoing": "\n".join(co_ongoing) if co_ongoing else "N/A",
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(base)

    data = load_yaml("cv_data.yaml")

    config = {}
    if os.path.exists("cancerfonden_config.yaml"):
        config = load_yaml("cancerfonden_config.yaml") or {}

    type_overrides = config.get("type_overrides", []) or []

    publications = data.get("publications", [])
    preprints = data.get("preprints", [])
    scholarly_debate = data.get("scholarly_debate", [])

    # ===================================================================
    # PUBLICATION COUNTS
    # ===================================================================
    total_original = 0
    total_first = 0
    total_last = 0
    total_review = 0

    for pub in publications:
        ptype = classify_publication(pub, type_overrides)
        if ptype == "original":
            total_original += 1
            if is_first_author(pub):
                total_first += 1
            if is_last_author(pub):
                total_last += 1
        elif ptype == "review":
            total_review += 1

    # ===================================================================
    # PUBLICATION LIST — classify 2018-2026
    # ===================================================================
    pubs_sorted = sorted(publications, key=lambda p: -(p.get("year", 0) or 0))

    originals_2018 = []
    reviews_2018 = []
    other_articles_2018 = []

    for pub in pubs_sorted:
        if not in_year_range(pub):
            continue
        ptype = classify_publication(pub, type_overrides)
        if ptype == "review":
            reviews_2018.append(pub)
        elif ptype == "other":
            other_articles_2018.append(pub)
        else:
            originals_2018.append(pub)

    # Note: "Other articles" (preprints, scholarly debate) excluded from Cancerfonden output

    # Selected 10 most important (by id)
    selected_ids = config.get("selected_publications", []) or []
    selected_pubs = []
    all_pubs_combined = publications + preprints
    for sel_id in selected_ids:
        for pub in all_pubs_combined:
            if pub.get("id") == sel_id:
                selected_pubs.append(pub)
                break
        else:
            print(f"WARNING: selected publication id '{sel_id}' not found in YAML")

    # ===================================================================
    # SUPERVISION
    # ===================================================================
    sup_fields = generate_supervision_fields(data)
    # Allow config overrides
    for key in ("main_completed", "main_ongoing", "co_completed", "co_ongoing"):
        config_key = f"supervision_{key}"
        if config.get(config_key):
            sup_fields[key] = config[config_key]

    # ===================================================================
    # OTHER SCIENTIFIC MERITS
    # ===================================================================
    other_merits = config.get("other_scientific_merits")
    if not other_merits:
        other_merits = generate_other_merits(data)

    # ===================================================================
    # BUILD WORD DOCUMENT (Publication list PDF)
    # ===================================================================
    doc = Document()
    set_style(doc)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{data.get('meta', {}).get('name', '')} – Publication List")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    # --- 10 most important ---
    add_heading(doc, "10 most important publications for this project", level=2)
    if selected_pubs:
        for i, pub in enumerate(selected_pubs, 1):
            add_reference_para(doc, pub, i)
    else:
        p = doc.add_paragraph()
        run = p.add_run("[Configure selected_publications in cancerfonden_config.yaml with 10 DOIs]")
        run.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)

    # --- Complete publication list 2018-2026 ---
    add_heading(doc, "Complete publication list 2018–2026", level=2)

    # Peer-reviewed original research articles
    if originals_2018:
        add_heading(doc, "Peer-reviewed original research articles", level=3)
        for i, pub in enumerate(originals_2018, 1):
            add_reference_para(doc, pub, i)

    # Review articles
    if reviews_2018:
        add_heading(doc, "Review articles", level=3)
        for i, pub in enumerate(reviews_2018, 1):
            add_reference_para(doc, pub, i)

    # Other articles section omitted per Cancerfonden requirements

    # Save docx
    os.makedirs("output", exist_ok=True)
    docx_path = os.path.join("output", "cancerfonden_publist.docx")
    doc.save(docx_path)
    print(f"Written: {docx_path}")

    # ===================================================================
    # TEXT FIELDS (for copy-paste into application form)
    # ===================================================================
    lines = []
    lines.append("=" * 70)
    lines.append("CANCERFONDEN CV FIELDS — Copy-paste into application form")
    lines.append("=" * 70)

    lines.append("")
    lines.append("--- PUBLICATION COUNTS ---")
    lines.append(f"Total number of original research articles (peer-reviewed): {total_original}")
    lines.append(f"Total number of original research articles as first author: {total_first}")
    lines.append(f"Total number of original research articles as last author:  {total_last}")
    lines.append(f"Total number of published review articles (peer-reviewed):  {total_review}")

    lines.append("")
    lines.append("--- OTHER SCIENTIFIC MERITS (max 1,500 characters) ---")
    lines.append(f"[{len(other_merits)} chars]")
    lines.append(other_merits)

    lines.append("")
    lines.append("--- SUPERVISION: Main supervisor, completed (max 1,000 chars) ---")
    lines.append(f"[{len(sup_fields['main_completed'])} chars]")
    lines.append(sup_fields["main_completed"])

    lines.append("")
    lines.append("--- SUPERVISION: Main supervisor, ongoing (max 1,000 chars) ---")
    lines.append(f"[{len(sup_fields['main_ongoing'])} chars]")
    lines.append(sup_fields["main_ongoing"])

    lines.append("")
    lines.append("--- SUPERVISION: Co-supervisor, completed (max 1,000 chars) ---")
    lines.append(f"[{len(sup_fields['co_completed'])} chars]")
    lines.append(sup_fields["co_completed"])

    lines.append("")
    lines.append("--- SUPERVISION: Co-supervisor, ongoing (max 1,000 chars) ---")
    lines.append(f"[{len(sup_fields['co_ongoing'])} chars]")
    lines.append(sup_fields["co_ongoing"])

    # Unpublished manuscripts
    manuscripts = config.get("unpublished_manuscripts", []) or []
    if manuscripts:
        lines.append("")
        lines.append("--- UNPUBLISHED MANUSCRIPTS (max 3) ---")
        for i, ms in enumerate(manuscripts, 1):
            lines.append(f"{i}. {ms.get('authors', '')}: {ms.get('title', '')}. {ms.get('status', '')}")

    lines.append("")
    lines.append("--- PUBLICATION LIST SUMMARY (2018-2026) ---")
    lines.append(f"Peer-reviewed original articles: {len(originals_2018)}")
    lines.append(f"Review articles:                 {len(reviews_2018)}")
    lines.append(f"TOTAL in list:                   {len(originals_2018) + len(reviews_2018)}")

    txt_output = "\n".join(lines)
    txt_path = os.path.join("output", "cancerfonden_cv_fields.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_output)
    print(f"Written: {txt_path}")
    print()
    print(txt_output)


if __name__ == "__main__":
    main()
